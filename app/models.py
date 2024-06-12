import logging
import os
import uuid

import pandas as pd
from django.conf import settings
from django.contrib.auth import get_user_model
from django.core.exceptions import ValidationError
from django.db import models
from django.db.models.signals import post_save
from django.dispatch import receiver
from docx import Document

User = get_user_model()

logger = logging.getLogger(__name__)


class Template(models.Model):
    id = models.UUIDField(default=uuid.uuid4, primary_key=True, editable=False)
    name = models.CharField(max_length=255, verbose_name='Name')
    file = models.FileField(upload_to='templates/', verbose_name='Template file')
    # owner = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name='Owner')
    # public = models.BooleanField(default=False, verbose_name='Public')
    uploaded_at = models.DateTimeField(auto_now_add=True, verbose_name='Uploaded at')

    def __str__(self):
        return f"Template - {self.name}"

    class Meta:
        ordering = ['-uploaded_at']
        verbose_name = 'Template'
        verbose_name_plural = 'Templates'

    def delete(self, using=None, keep_parents=False):
        try:
            if self.file:
                os.remove(self.file.path)
        except Exception as e:
            logger.error(f"Error while deleting file {self.file.path}: {e}")
        super().delete(using, keep_parents)


class DataSource(models.Model):
    id = models.UUIDField(default=uuid.uuid4, primary_key=True, editable=False)
    name = models.CharField(max_length=255, verbose_name='Name')
    file = models.FileField(upload_to='datasources/', verbose_name='Datasource file')
    # owner = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name='Owner')
    # public = models.BooleanField(default=False, verbose_name='Public')
    uploaded_at = models.DateTimeField(auto_now_add=True, verbose_name='Uploaded at')

    def __str__(self):
        return f"Datasource - {self.name}"

    class Meta:
        ordering = ['-uploaded_at']
        verbose_name = 'Datasource'
        verbose_name_plural = 'Datasources'

    def delete(self, using=None, keep_parents=False):
        try:
            if self.file:
                os.remove(self.file.path)
        except Exception as e:
            logger.error(f"Error while deleting file {self.file.path}: {e}")
        super().delete(using, keep_parents)


class RenderedFile(models.Model):
    id = models.UUIDField(default=uuid.uuid4, primary_key=True, editable=False)
    file = models.FileField(upload_to='renders/', verbose_name='Rendered file')
    datasource = models.ForeignKey(DataSource, on_delete=models.SET_NULL, null=True)
    template = models.ForeignKey(Template, on_delete=models.SET_NULL, null=True)
    sheet = models.CharField(max_length=255, default='Sheet2')
    # owner = models.ForeignKey(User, on_delete=models.CASCADE)
    # public = models.BooleanField(default=False)
    uploaded_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        datasource = self.datasource.name if self.datasource else 'No datasource'
        template = self.template.name if self.template else 'No template'
        return f"RenderedFile {datasource} => {template}"

    class Meta:
        ordering = ['-uploaded_at']
        verbose_name = 'Rendered File'
        verbose_name_plural = 'Rendered Files'

    def delete(self, using=None, keep_parents=False):
        try:
            if self.file:
                os.remove(self.file.path)
        except Exception as e:
            logger.error(f"Error while deleting file {self.file.path}: {e}")
        super().delete(using, keep_parents)


### GOOD VERSION

# @receiver(post_save, sender=RenderedFile)
# def process_file(sender, instance, **kwargs):
#     if not validate_instance(instance):
#         return
#
#     df = read_excel_file(instance)
#     if df is None:
#         return
#
#     doc = process_word_document(instance, df)
#     replace_all_placeholders(doc, df)
#     save_document(instance, doc)
#
#
# def validate_instance(instance):
#     if instance.file:
#         return False
#     if not instance.datasource or not instance.template:
#         instance.delete()
#         return False
#     if not instance.sheet:
#         instance.sheet = 'Sheet1'
#     return True
#
#
# def read_excel_file(instance):
#     try:
#         return pd.read_excel(instance.datasource.file.path, sheet_name=instance.sheet)
#     except ValueError:
#         instance.delete()
#         raise ValidationError("Sheet not found")
#
#
# def process_word_document(instance, df):
#     doc = Document(instance.template.file.path)
#     for i, row in df.iterrows():
#         process_paragraphs(doc.paragraphs, row)
#         process_tables(doc.tables, row)
#     return doc
#
#
# def process_paragraphs(paragraphs, row):
#     for paragraph in paragraphs:
#         for key in row.keys():
#             replace_placeholder(paragraph, key, row[key])
#
#
# def process_tables(tables, row):
#     for table in tables:
#         template_tag_found = False
#         for row_ in table.rows:
#             for cell in row_.cells:
#                 process_paragraphs(cell.paragraphs, row)
#         if template_tag_found:
#             copy_and_replace_last_row(table, row)
#             template_tag_found = False
#
#
# def replace_placeholder(paragraph, key, value):
#     if f'{{{{{key}}}}}' in paragraph.text:
#         if hasattr(value, 'strftime'):
#             value = value.strftime('%d.%m.%Y')
#         for run in paragraph.runs:
#             run.text = run.text.replace(f'{{{{{key}}}}}', str(value))
#         # paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))
#
#
# def copy_and_replace_last_row(table, row):
#     last_row = table.rows[-1]
#     new_row = table.add_row()
#     for cell, last_cell in zip(new_row.cells, last_row.cells):
#         for paragraph, last_paragraph in zip(cell.paragraphs, last_cell.paragraphs):
#             paragraph.text = last_paragraph.text
#             for key in row.keys():
#                 replace_placeholder(paragraph, key, row[key])
#
#
# def replace_all_placeholders(doc, df):
#     for paragraph in doc.paragraphs:
#         replace_all_in_paragraph(paragraph, df)
#     for table in doc.tables:
#         for row_ in table.rows:
#             for cell in row_.cells:
#                 replace_all_in_paragraph(cell, df)
#
#
# def replace_all_in_paragraph(paragraph, df):
#     for key in df.columns:
#         if f'{{{{{key}|all}}}}' in paragraph.text:
#             if hasattr(df.iloc[0][key], 'strftime'):
#                 all_values = ', '.join([date.strftime('%d.%m.%Y') for date in df[key] if not pd.isna(date)])
#             else:
#                 all_values = ', '.join([str(value) for value in df[key] if not pd.isna(value)])
#             for run in paragraph.runs:
#                 run.text = run.text.replace(f'{{{{{key}|all}}}}', all_values)
#
#
# def save_document(instance, doc):
#     doc_path = os.path.join(settings.MEDIA_ROOT, 'renders', f'{instance.id}.docx')
#     doc.save(doc_path)
#     instance.file = doc_path
#     instance.save()
#     logger.info(f"Rendered file {instance.id} saved at {doc_path}")
#     logger.info(f"Rendered file {instance.id} saved in database")

@receiver(post_save, sender=RenderedFile)
def process_file(sender, instance, **kwargs):
    if not validate_instance(instance):
        return

    df = read_excel_file(instance)
    if df is None:
        return

    doc = process_word_document(instance, df)
    replace_all_placeholders(doc, df)
    save_document(instance, doc)


def validate_instance(instance):
    if instance.file:
        return False
    if not instance.datasource or not instance.template:
        instance.delete()
        return False
    if not instance.sheet:
        instance.sheet = 'Sheet1'
    return True


def read_excel_file(instance):
    try:
        return pd.read_excel(instance.datasource.file.path, sheet_name=instance.sheet)
    except ValueError:
        instance.delete()
        raise ValidationError("Sheet not found")


def process_word_document(instance, df):
    doc = Document(instance.template.file.path)
    for i, row in df.iterrows():
        process_paragraphs(doc.paragraphs, row)
        process_tables(doc.tables, row)
    return doc


def process_paragraphs(paragraphs, row):
    for paragraph in paragraphs:
        for key in row.keys():
            replace_placeholder(paragraph, key, row[key])


def process_tables(tables, row):
    for table in tables:
        for row_ in table.rows:
            for cell in row_.cells:
                process_paragraphs(cell.paragraphs, row)


def replace_placeholder(paragraph, key, value):
    if f'{{{{{key}}}}}' in paragraph.text:
        if hasattr(value, 'strftime'):
            value = value.strftime('%d.%m.%Y')
        for run in paragraph.runs:
            run.text = run.text.replace(f'{{{{{key}}}}}', str(value))


def replace_all_placeholders(doc, df):
    for paragraph in doc.paragraphs:
        replace_all_in_paragraph(paragraph, df)
    for table in doc.tables:
        for row_ in table.rows:
            for cell in row_.cells:
                replace_all_in_paragraph(cell, df)


def replace_all_in_paragraph(paragraph, df):
    for key in df.columns:
        if f'{{{{{key}|all}}}}' in paragraph.text:
            if hasattr(df.iloc[0][key], 'strftime'):
                all_values = ', '.join([date.strftime('%d.%m.%Y') for date in df[key] if not pd.isna(date)])
            else:
                all_values = ', '.join([str(value) for value in df[key] if not pd.isna(value)])
            for run in paragraph.runs:
                run.text = run.text.replace(f'{{{{{key}|all}}}}', all_values)


def save_document(instance, doc):
    doc_path = os.path.join(settings.MEDIA_ROOT, 'renders', f'{instance.id}.docx')
    doc.save(doc_path)
    instance.file = doc_path
    instance.save()
    logger.info(f"Rendered file {instance.id} saved at {doc_path}")
    logger.info(f"Rendered file {instance.id} saved in database")
